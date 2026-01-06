# T5_report_word__v5.py
# 變更：新增檔名安全化處理（避免 Windows 命名失敗）+ 長路徑保存
# 其他維持：EX 插入點、手動紅色連編、網址必佔號+超連結、V2 裁白邊與容錯、V4 書籤

from typing import List, Dict, Optional
from pathlib import Path
import os, re, datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ----------------------- 圖片裁白邊（沿用 v2） -----------------------
def _autocrop_lr_whitespace(src_path: str, out_suffix: str = "_crop.png") -> str:
    try:
        im_color = Image.open(src_path)
        im_gray = im_color.convert("L")
        W, H = im_gray.size
        arr = im_gray.load()

        threshold = 245
        min_margin = 30
        max_cut_ratio = 0.2

        left = 0
        for x in range(W):
            col_brightness = sum(arr[x, y] for y in range(H)) / H
            if col_brightness < threshold:
                left = max(0, x - min_margin)
                break

        right = W
        for x in range(W - 1, -1, -1):
            col_brightness = sum(arr[x, y] for y in range(H)) / H
            if col_brightness < threshold:
                right = min(W, x + min_margin)
                break

        max_cut = int(W * max_cut_ratio)
        left = min(left, max_cut)
        right = max(right, W - max_cut)

        if right <= left:
            return src_path

        im2 = im_color.crop((left, 0, right, H))
        out_path = src_path[:-4] + out_suffix
        im2.save(out_path)
        return out_path
    except Exception:
        return src_path

# ----------------------- Word 初始化（沿用 v2） -----------------------
def _init_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "新細明體"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    return doc

def _safe_text(x: Optional[str], default: str = "") -> str:
    s = (x or "").strip()
    return s if s else default

def _usable_width_emu(doc: Document) -> int:
    sec = doc.sections[-1]
    return int(sec.page_width - sec.left_margin - sec.right_margin)

def _is_existing_file(p: str) -> bool:
    try:
        return bool(p) and os.path.exists(p)
    except Exception:
        return False

# ----------------------- 超連結 helper -----------------------
def _add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as oxml_qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(oxml_qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

# ----------------------- 書籤 helper -----------------------
def _add_bookmark(paragraph, name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as oxml_qn
    b_id = abs(hash(name)) % (2**15)
    start = OxmlElement("w:bookmarkStart")
    start.set(oxml_qn("w:id"), str(b_id))
    start.set(oxml_qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(oxml_qn("w:id"), str(b_id))
    p = paragraph._p
    p.insert(0, start)
    p.append(end)

# ----------------------- 找「正本：」後的插入點 -----------------------
def _find_insert_index_after_zhengben(lines: List[str]) -> Optional[int]:
    if not lines:
        return None
    for i, line in enumerate(lines):
        if "正本：" in (line or ""):
            return min(i + 1, len(lines))
    return None

# ----------------------- Windows 長路徑（新增） -----------------------
def _as_longpath(p: str) -> str:
    """將路徑轉為 Windows 長路徑前綴 \\?\ ，避免 260 字元上限。非 Windows 原樣返回。"""
    try:
        import sys
        ap = os.path.abspath(p)
        if sys.platform.startswith("win") and not ap.startswith("\\\\?\\") and len(ap) >= 240:
            return "\\\\?\\" + ap
        return ap
    except Exception:
        return p

# ----------------------- 檔名安全化（新增） -----------------------
def _sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", name or "")
    name = name.strip(" .")
    if len(name) > 120:
        name = name[:120]
    # 確保有 .docx 副檔名（避免截短後遺失）
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name or datetime.datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx"

# ----------------------- 單一商品條目 -----------------------
def _insert_one_item_block(doc: Document, r: Dict, url_text: str, usable_w: int, item_no: int, add_bookmark: bool) -> None:
    title = _safe_text(r.get("api_title")) or _safe_text(r.get("name")) or "商品名稱未找到"
    bsmi = _safe_text(r.get("bsmi"), "查無")
    model_no = _safe_text(r.get("model_no"), "查無")
    seller_acc = _safe_text(r.get("seller_account"))

    # 🔴 商品標題（唯一進導覽的地方）
    p_title = doc.add_paragraph()
    p_title.style = "Heading 3"   # ← 關鍵：讓這一整行進導覽

    run_no = p_title.add_run(f"{item_no}. ")
    run_no.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    run_title = p_title.add_run(title)
    run_title.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    if add_bookmark:
        _add_bookmark(p_title, f"item-{item_no:03d}")

    # 以下全部不進導覽（保持一般段落）
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"(商品檢驗標識：{bsmi}、型號：{model_no})")
    run_sub.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    if seller_acc:
        doc.add_paragraph(f"賣家帳號：{seller_acc}")

    p_url = doc.add_paragraph("網址：")
    if url_text and url_text.lower().startswith("http"):
        _add_hyperlink(p_url, url_text, url_text)
    else:
        p_url.add_run("查無")

    targets = []
    pngs = r.get("pngs") or []
    if pngs:
        targets = [p for p in pngs if _is_existing_file(p)]
    else:
        p = _safe_text(r.get("png"))
        if _is_existing_file(p):
            targets = [p]

    for p in targets:␊
        try:␊
            p2 = _autocrop_lr_whitespace(p)␊
            doc.add_picture(p2, width=usable_w)␊
            para = doc.paragraphs[-1]␊
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = para.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)
        except Exception as e:
            doc.add_paragraph(f"[圖片插入失敗: {os.path.basename(p)} - {e}]")

    desc_imgs = r.get("desc_imgs") or []
    if desc_imgs:
        doc.add_paragraph("描述圖片（供人工審核）：")
        for p in desc_imgs:
            if not _is_existing_file(p):
                continue
            try:
                p2 = _autocrop_lr_whitespace(p)
                doc.add_picture(p2, width=usable_w)
                para = doc.paragraphs[-1]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fmt = para.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(6)
            except Exception as e:
                doc.add_paragraph(f"[圖片插入失敗: {os.path.basename(p)} - {e}]")
        doc.add_paragraph("")
# ----------------------- 主 API -----------------------
def render_word(segments: List[Dict], out_docx: str) -> str:
    out_path = str(Path(out_docx))
    safe_name = _sanitize_filename(Path(out_path).name)
    out_path = str(Path(out_path).with_name(safe_name))
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    doc = _init_doc()
    usable_w = _usable_width_emu(doc)
    item_no = 1

    for sidx, seg in enumerate(segments or []):
        lines: List[str] = seg.get("lines") or []
        urls: List[str] = seg.get("urls") or []
        results: List[Dict] = seg.get("results") or []
        insert_at = _find_insert_index_after_zhengben(lines)

        def _insert_items_here():
            nonlocal item_no
            n_items = max(len(urls), len(results))
            for i in range(n_items):
                r = results[i] if i < len(results) else {}
                u = urls[i] if i < len(urls) else ""
                _insert_one_item_block(doc, r, u, usable_w, item_no, True)
                item_no += 1

        if insert_at is not None:
            before = lines[:insert_at]
            after = lines[insert_at:]
            for line in before: doc.add_paragraph(_safe_text(line))
            _insert_items_here()
            for line in after: doc.add_paragraph(_safe_text(line))
        else:
            for line in lines: doc.add_paragraph(_safe_text(line))
            _insert_items_here()

        if sidx != len(segments) - 1:
            doc.add_paragraph("")

    try:
        # 長路徑保護
        doc.save(_as_longpath(out_path))
    except Exception:
        alt = str(Path(out_path).with_name(Path(out_path).stem + "_1.docx"))
        try:
            doc.save(_as_longpath(alt))
            out_path = alt
        except Exception:
            pass
    return out_path

